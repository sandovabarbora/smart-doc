import os
import hashlib
from typing import List, Dict, Any, Tuple
from pathlib import Path
import asyncio
from concurrent.futures import ThreadPoolExecutor

import PyPDF2
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

from ..core.config import settings
from ..models.document import Document, DocumentChunk
from sqlalchemy.orm import Session

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""]
        )
        self.executor = ThreadPoolExecutor(max_workers=4)
    
    async def process_document(self, file_path: str, document: Document, db: Session) -> List[DocumentChunk]:
        try:
            content = await self._extract_content(file_path, document.content_type)
            chunks = await self._create_chunks(content, document.filename)
            
            db_chunks = []
            for i, chunk in enumerate(chunks):
                db_chunk = DocumentChunk(
                    document_id=document.id,
                    chunk_index=i,
                    content=chunk.page_content,
                    metadata=chunk.metadata
                )
                db_chunks.append(db_chunk)
                db.add(db_chunk)
            
            document.processed = True
            db.commit()
            return db_chunks
            
        except Exception as e:
            document.processing_error = str(e)
            document.processed = False
            db.commit()
            raise
    
    async def _extract_content(self, file_path: str, content_type: str) -> str:
        loop = asyncio.get_event_loop()
        
        if content_type == "application/pdf":
            return await loop.run_in_executor(self.executor, self._extract_pdf, file_path)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return await loop.run_in_executor(self.executor, self._extract_docx, file_path)
        elif content_type in ["text/plain", "text/markdown"]:
            return await loop.run_in_executor(self.executor, self._extract_text, file_path)
        else:
            raise ValueError(f"Unsupported content type: {content_type}")
    
    def _extract_pdf(self, file_path: str) -> str:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            return text
    
    def _extract_docx(self, file_path: str) -> str:
        doc = DocxDocument(file_path)
        paragraphs = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    paragraphs.append(f"TABLE: {row_text}")
        
        return "\n\n".join(paragraphs)
    
    def _extract_text(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    async def _create_chunks(self, content: str, filename: str) -> List[LangchainDocument]:
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            self.executor, 
            self._chunk_content, 
            content, 
            filename
        )
    
    def _chunk_content(self, content: str, filename: str) -> List[LangchainDocument]:
        content = self._preprocess_content(content)
        texts = self.text_splitter.split_text(content)
        
        documents = []
        for i, text in enumerate(texts):
            metadata = {
                "source": filename,
                "chunk_id": i,
                "chunk_size": len(text),
                "content_hash": hashlib.md5(text.encode()).hexdigest()
            }
            metadata.update(self._extract_metadata(text))
            
            documents.append(LangchainDocument(
                page_content=text,
                metadata=metadata
            ))
        
        return documents
    
    def _preprocess_content(self, content: str) -> str:
        lines = content.split('\n')
        processed_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                line = ' '.join(line.split())
                if self._is_header(line):
                    line = f"\n\n## {line}\n"
                processed_lines.append(line)
        
        return ' '.join(processed_lines)
    
    def _is_header(self, line: str) -> bool:
        return (
            line.isupper() and len(line) < 100 or
            any(line.startswith(prefix) for prefix in ['Chapter', 'Section', 'Part']) or
            line.endswith(':') and len(line) < 100
        )
    
    def _extract_metadata(self, text: str) -> Dict[str, Any]:
        metadata = {}
        
        if any(keyword in text.lower() for keyword in ['table', 'row', 'column', '|']):
            metadata['content_type'] = 'table'
        elif any(keyword in text.lower() for keyword in ['figure', 'chart', 'graph', 'image']):
            metadata['content_type'] = 'figure_reference'
        elif text.count('\n') > 5:
            metadata['content_type'] = 'structured'
        else:
            metadata['content_type'] = 'paragraph'
        
        words = text.lower().split()
        if len(words) > 0:
            metadata['word_count'] = len(words)
            metadata['starts_with'] = words[0] if words else ""
        
        return metadata
    
    async def cleanup_temp_file(self, file_path: str):
        try:
            os.remove(file_path)
        except OSError:
            pass
